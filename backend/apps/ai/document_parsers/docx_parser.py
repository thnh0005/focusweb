import zipfile
from io import BytesIO
from xml.etree import ElementTree

from .base import BaseDocumentParser, ExtractionResult
from .exceptions import FileTypeMismatchError, InvalidDOCXError


WORD_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


class DOCXParser(BaseDocumentParser):
    file_type = "docx"
    extensions = {".docx"}
    mime_types = {
        "",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/octet-stream",
    }

    def validate_signature(self, content):
        if not content.startswith(b"PK"):
            raise FileTypeMismatchError()
        try:
            with zipfile.ZipFile(BytesIO(content)) as archive:
                names = set(archive.namelist())
                if "[Content_Types].xml" not in names or "word/document.xml" not in names:
                    raise InvalidDOCXError()
        except zipfile.BadZipFile as exc:
            raise InvalidDOCXError() from exc
        return True

    def extract(self, content):
        self.validate_signature(content)
        try:
            return self.extract_with_python_docx(content)
        except ImportError:
            return self.extract_with_stdlib(content)

    def extract_with_python_docx(self, content):
        from docx import Document

        try:
            document = Document(BytesIO(content))
        except Exception as exc:
            raise InvalidDOCXError() from exc

        blocks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    blocks.append(" | ".join(cells))
        return self.result_from_blocks(blocks)

    def extract_with_stdlib(self, content):
        try:
            with zipfile.ZipFile(BytesIO(content)) as archive:
                xml = archive.read("word/document.xml")
        except (KeyError, zipfile.BadZipFile) as exc:
            raise InvalidDOCXError() from exc

        try:
            root = ElementTree.fromstring(xml)
        except ElementTree.ParseError as exc:
            raise InvalidDOCXError() from exc

        blocks = []
        for child in root.findall(".//w:body/*", WORD_NS):
            if child.tag.endswith("}p"):
                text = self.paragraph_text(child)
                if text:
                    blocks.append(text)
            elif child.tag.endswith("}tbl"):
                for row in child.findall(".//w:tr", WORD_NS):
                    cells = [
                        self.paragraph_text(cell)
                        for cell in row.findall(".//w:tc", WORD_NS)
                    ]
                    cells = [cell for cell in cells if cell]
                    if cells:
                        blocks.append(" | ".join(cells))

        return self.result_from_blocks(blocks)

    def result_from_blocks(self, blocks):
        text = "\n".join(blocks)
        return ExtractionResult(
            text=text,
            page_count=None,
            page_map=[
                {
                    "section": 1,
                    "type": "paragraph_group",
                    "start_char": 0,
                    "end_char": len(text),
                    "character_count": len(text),
                }
            ]
            if text
            else [],
            detected_file_type=self.file_type,
        )

    def paragraph_text(self, node):
        text = "".join(item.text or "" for item in node.findall(".//w:t", WORD_NS))
        return text.strip()
